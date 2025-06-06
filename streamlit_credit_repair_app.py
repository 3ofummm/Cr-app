
import streamlit as st
from docx import Document
from datetime import date
from io import BytesIO

st.title("Automated Credit Dispute Letter Generator")

st.write("This tool generates a detailed credit dispute letter using dummy data. In the final version, you'll be able to upload or enter your real credit report details.")

name = st.text_input("Your Full Name", "John Doe")
address = st.text_input("Your Address", "1234 Main Street")
city_state_zip = st.text_input("City, State ZIP", "Hometown, ST 12345")
today_date = date.today().strftime("%B %d, %Y")

if st.button("Generate Dispute Letter"):
    doc = Document()
    doc.add_paragraph(name)
    doc.add_paragraph(address)
    doc.add_paragraph(city_state_zip)
    doc.add_paragraph(today_date)
    doc.add_paragraph("")
    doc.add_paragraph("Experian")
    doc.add_paragraph("P.O. Box 4500")
    doc.add_paragraph("Allen, TX 75013")
    doc.add_paragraph("")
    doc.add_paragraph("RE: Request for Investigation under FCRA §609 and §611")
    doc.add_paragraph("To Whom It May Concern,")
    doc.add_paragraph("")
    doc.add_paragraph("I am writing to formally dispute the accuracy of the following information listed on my credit report. In accordance with the Fair Credit Reporting Act (§609 and §611), I am requesting an investigation and removal of the following inaccurate and unverifiable information:")
    doc.add_paragraph("1. Account: Capital One / #1234\n   - Reason: This account shows a 30-day late payment from 03/2020, which is inaccurate. I have never made a late payment on this account. Please verify or remove this entry.")
    doc.add_paragraph("2. Account: Midland Credit Management / #5678\n   - Reason: This collection account is duplicated from a previously listed item and is over 7 years old. Please delete.")
    doc.add_paragraph("3. Personal Information\n   - Incorrect name: 'Jon Smith' – I have never used this name.\n   - Incorrect employer: 'ABC Construction' – I’ve never worked there.")
    doc.add_paragraph("")
    doc.add_paragraph("Please provide a copy of the results of your investigation, as required by FCRA. I have enclosed a copy of my ID and utility bill for verification purposes.")
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(name)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Dispute Letter (DOCX)",
        data=buffer,
        file_name="Dispute_Letter.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
